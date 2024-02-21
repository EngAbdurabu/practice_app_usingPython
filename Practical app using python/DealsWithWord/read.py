import docx, ReadDocx
from pathlib import Path

doc = docx.Document(Path.home()/Path("OneDrive", "سطح المكتب", "python_practical",
                                     "academy_1.docx"))

print(len(doc.paragraphs))

print(doc.paragraphs[0].text)
print("=" * 20)
print(doc.paragraphs[1].text)
print("=" * 20)
print(doc.paragraphs[2].text)
print("=" * 20)

print(len(doc.paragraphs[2].runs))
print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)

print("*" * 50)
# read whole file




print(ReadDocx.GetText(Path.home()/Path("OneDrive", "سطح المكتب", "python_practical",
                                     "academy_1.docx")))

