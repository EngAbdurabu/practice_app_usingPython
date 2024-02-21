import docx
from pathlib import Path
# for control alignment
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


mydoc = docx.Document()
# to add paragraph
mydoc.add_paragraph("Hi I'm Engineer  Abdurabu Mohammed ")
mydoc.add_paragraph("I got the grad in this year  ")

# اضافة نص ب اللغة العربية
paragraph = mydoc.add_paragraph("يجب عليك الاستمرار في عمل الشيء حتى الاتقان")
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

third_paragraph = mydoc.add_paragraph("This is third paragraph in this file, ")
# to add text to the same paragraph without break line
third_paragraph.add_run("this is the end section of third paragraph")

# add header
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_heading("This is level 2 heading", 1)
mydoc.add_heading("This is level 3 heading", 2)
mydoc.add_heading("This is level 4 heading", 3)
mydoc.add_heading("This is level 5 heading", 4)

# style
print(mydoc.paragraphs[0].style)
print(mydoc.paragraphs[4].style)
print(mydoc.paragraphs[4].text)

print(mydoc.paragraphs[5].style)
print(mydoc.paragraphs[5].text)

# change style
mydoc.paragraphs[0].style = mydoc.styles["Title"]
mydoc.paragraphs[3].style = mydoc.styles["Heading 2"]

# remove style
mydoc.paragraphs[3].style.delete()
# another method to change style
mydoc.add_paragraph("this is last paragraph", "Title")

# add picture
mydoc.add_picture(str(Path.home()/Path("OneDrive", "سطح المكتب", "python_practical",
                                     "code.png")), width=docx.shared.Inches(5),
                                                height=docx.shared.Inches(3))



# save document
mydoc.save(Path.home()/Path("OneDrive", "سطح المكتب", "python_practical",
                                     "Write.docx"))